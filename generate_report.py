import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def generate_student_report(output_folder, report_path="Student_Report.docx"):
    """
    Generate a grid-format student image report.
    
    Args:
        output_folder (str): Path to KinderSort output folder (contains student subfolders).
        report_path (str): Path to save the output Word report.
    """
    # Create a new Word document
    doc = Document()
    
    # Set title
    title = doc.add_heading('Student Photo Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Get all student subfolders (exclude _unmatched and log files)
    output_path = Path(output_folder)
    student_folders = [f for f in output_path.iterdir() if f.is_dir() and f.name != "_unmatched"]
    
    if not student_folders:
        doc.add_paragraph("No student folders found.")
        doc.save(report_path)
        return
    
    # Create a table for each student (3 columns for grid layout)
    for student_folder in student_folders:
        student_name = student_folder.name
        doc.add_heading(f'Student: {student_name}', level=2)
        
        # Get all image files in this student's folder
        image_files = [f for f in student_folder.iterdir() 
                       if f.suffix.lower() in ['.jpg', '.jpeg', '.png', '.gif']]
        
        if not image_files:
            doc.add_paragraph("No images found for this student.")
            continue
        
        # Create table: rows calculated dynamically, 3 columns per row
        cols = 3
        rows = (len(image_files) + cols - 1) // cols  # Round up
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Fill table cells
        for i, img_path in enumerate(image_files):
            row_idx = i // cols
            col_idx = i % cols
            cell = table.cell(row_idx, col_idx)
            
            # Insert image into cell
            try:
                # Set image width to 1.5 inches for neat grid
                run = cell.paragraphs[0].add_run()
                run.add_picture(str(img_path), width=Inches(1.5))
                # Add filename below image (optional)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                cell.text = f"[Error: {os.path.basename(img_path)}]"
        
        doc.add_paragraph()  # Add blank line between students
    
    # Save document
    doc.save(report_path)
    print(f"Report saved to: {report_path}")

if __name__ == "__main__":
    # Example usage: assuming output folder is "./output"
    generate_student_report("./output", "Student_Report.docx")